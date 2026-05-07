from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "Alix-Vintage_Project-Documentation_Updated.md"
DOCX_PATH = ROOT / "docs" / "Alix-Vintage_Project-Documentation_Updated.docx"


H1_SET = {
    "BUSINESS REQUIREMENTS",
    "USER REQUIREMENTS",
    "SYSTEM REQUIREMENTS SPECIFICATION",
    "ACTIVITY DIAGRAM",
    "USE CASE DIAGRAM",
    "CONTEXT FLOW DIAGRAM",
    "LEVEL-0 DATA FLOW DIAGRAM",
    "SOFTWARE PROCESS MODEL (METHODOLOGY)",
}

H2_SET = {
    "Table of Contents",
    "Business Profile",
    "Problem Statement",
    "Project Constraints",
    "Project Objectives",
    "Project Scope Statement",
    "Business Process Analysis",
    "Stakeholder Analysis",
    "Stakeholders Needs",
    "Current Process Issues",
    "Proposed Process Improvements",
    "In Scope",
    "Out of Scope",
    "Table 1. User Requirements(Admin - Internal User)",
    "Table 2. User Requirements(Customer - Verified User)",
    "Table 3. Functional Requirements (Admin)",
    "Table 4. Functional Requirements (Verified Customer)",
    "Table 5. Non-functional Requirements",
    "Verified Customer Workflow",
    "Admin Workflow",
    "Verified Customer Use Cases",
    "Administrator Use Cases",
    "Verified Customer Information Flows",
    "Administrator Information Flows",
    "Phases and Activities",
    "1. Requirement Analysis",
    "2. System and Software Design",
    "3. Implementation / Coding",
    "4. Testing",
    "5. Deployment",
    "6. Maintenance",
}


TABLE_1 = [
    ("UR A1", "Admin Authentication", "a. Log in securely using encrypted credentials to access the backend management system."),
    (
        "UR A2",
        "Order Management",
        "a. Review submitted rosters and individual designs to ensure they are production-ready.\n"
        "b. Review custom upload requests and provide a manual price quote and shipping fee when applicable.\n"
        "c. Review fixed catalog orders and confirm material availability before fulfillment proceeds.\n"
        "d. Upload digital design mockups (Proofs) for customer review.\n"
        "e. Verify uploaded receipts for initial/partial payments and (when applicable) final payments.\n"
        "f. Manually transition orders through the system-supported lifecycle stages (pending, paid, proofing, processing, awaiting_final_payment, ready_to_ship, shipped, completed, cancelled).\n"
        "g. Input the J&T tracking number during the On Transit (shipped) phase.",
    ),
    ("UR A3", "Item Management", "a. Update the website showcase by adding new designs or toggling items as Out of Stock."),
    ("UR A4", "Data Archiving", "Access a master list of historical orders to support auditing and reporting."),
]

TABLE_2 = [
    (
        "UR C1",
        "Identity Verification",
        "a. Receive and enter a 6-digit OTP via email during registration to activate the account.\n"
        "b. Securely log in once the email has been verified.",
    ),
    (
        "UR C2",
        "Personalization",
        "a. Select apparel types and input bulk roster details (Names, Numbers, Sizes) via a dynamic grid (group orders).\n"
        "b. Upload design files or reference images with specific details for custom-made orders.",
    ),
    (
        "UR C3",
        "Transaction",
        "a. Choose the preferred payment method available (e.g., GCash or COD based on policy).\n"
        "b. Review and accept the Admin’s quote when manual quoting applies.\n"
        "c. Approve the Admin’s design proof or request revision with specific feedback.\n"
        "d. Upload payment receipts for verification when required.",
    ),
    (
        "UR C4",
        "Order Tracking",
        "a. Access a dashboard to view real-time order status.\n"
        "b. Copy the J&T tracking number and redirect to the official J&T tracking site.",
    ),
    (
        "UR C5",
        "Account History",
        "a. View a chronological list of previous orders, including status, totals, and final approved designs.",
    ),
]

TABLE_3 = [
    ("FR A1", "Product Management", "a. Admin must be able to add, update, and delete apparel listings, including images and pricing."),
    (
        "FR A2",
        "Order Review",
        "a. Admin can view roster details or individual design details for pending orders/requests.\n"
        "b. Admin must confirm stock availability and accept the request when required.\n"
        "c. Admin must input or adjust base price and shipping fee where applicable.\n"
        "d. Admin must be able to upload layout mockups for customer approval (proofing).",
    ),
    (
        "FR A3",
        "Payment Verification",
        "a. Admin must verify uploaded receipts based on workflow stage.\n"
        "b. Production should not proceed unless required payments are verified according to policy.",
    ),
    (
        "FR A4",
        "Fulfillment",
        "a. Admin must transition orders to Ready to Ship once production is complete.\n"
        "b. Admin must input a J&T tracking number and transition the status to On Transit once handed to the courier.",
    ),
]

TABLE_4 = [
    (
        "FR C1",
        "Email Verification",
        "a. Users must receive and enter a 6-digit confirmation code via email to verify their account during registration.\n"
        "b. Users must be able to request/resend a verification code.",
    ),
    (
        "FR C2",
        "Product Pick",
        "a. Users can select designs and choose group/individual paths where applicable.\n"
        "b. Users must be able to upload custom files or reference pictures if selecting the Upload Own Design path.",
    ),
    (
        "FR C3",
        "Roster Management",
        "a. Users can input multiple names, numbers, and sizes via a responsive grid and upload a logo for the batch.",
    ),
    (
        "FR C4",
        "Checkout & Obligation",
        "a. Users must agree to the Terms and Conditions before submitting the order.\n"
        "b. Users must upload required receipts and wait for Admin verification when applicable.\n"
        "c. Proofing Gate: Users must click Approve or Request Revision on the layout proof before production begins.",
    ),
    (
        "FR C5",
        "Status Monitoring",
        "a. The system must provide a My Orders portal showing system-defined stages.\n"
        "b. The portal must feature a copy button for tracking number and a link to the courier tracking website.",
    ),
    (
        "FR C6",
        "History Retrieval",
        "a. The system must query the database for records matching the user and display them in a Past Orders view.",
    ),
]

TABLE_5 = [
    ("NFR 1", "Performance", "a. Website pages must load in under 3 seconds."),
    ("NFR 2", "Usability", "a. The website must follow a Mobile-First Design approach."),
    (
        "NFR 3",
        "Security",
        "a. Passwords must be hashed in the database.\n"
        "b. Access to receipts, rosters, and custom design files must be restricted.",
    ),
    ("NFR 4", "Compatibility", "a. The website must be responsive on Android, iOS, and modern desktop browsers."),
    ("NFR 5", "Reliability", "a. The system must maintain proof version history and revision notes for each order item."),
]


def _set_default_font(document: Document, font_name: str = "Times New Roman", size_pt: int = 12) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = Pt(size_pt)


def _add_centered_cover(document: Document, cover_lines: list[str]) -> None:
    for line in cover_lines:
        p = document.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
    document.add_page_break()


def _add_heading(document: Document, text: str, level: int) -> None:
    document.add_heading(text, level=level)


def _add_table(document: Document, rows: list[tuple[str, str, str]]) -> None:
    table = document.add_table(rows=len(rows) + 1, cols=3)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Requirement No."
    hdr[1].text = "Category"
    hdr[2].text = "Description"

    for idx, (req_no, cat, desc) in enumerate(rows, start=1):
        cells = table.rows[idx].cells
        cells[0].text = req_no
        cells[1].text = cat
        cells[2].text = desc

    # Slightly widen description column on typical letter/A4.
    for row in table.rows:
        row.cells[0].width = Inches(1.0)
        row.cells[1].width = Inches(1.6)
        row.cells[2].width = Inches(4.8)

    document.add_paragraph("")


def build_docx(md_text: str) -> Document:
    lines = [ln.rstrip() for ln in md_text.splitlines()]

    # Cover page: everything before "Table of Contents".
    cover_lines: list[str] = []
    body_start = 0
    for i, ln in enumerate(lines):
        if ln.strip() == "Table of Contents":
            body_start = i
            break
        if ln.strip() == "":
            continue
        cover_lines.append(ln.strip())

    doc = Document()
    _set_default_font(doc)

    if cover_lines:
        _add_centered_cover(doc, cover_lines)

    i = body_start
    while i < len(lines):
        ln = lines[i].strip()

        if ln == "":
            i += 1
            continue

        # Tables
        if ln == "Table 1. User Requirements(Admin - Internal User)":
            _add_heading(doc, ln, 2)
            _add_table(doc, TABLE_1)
            i += 1
            while i < len(lines) and lines[i].strip() != "Table 2. User Requirements(Customer - Verified User)":
                i += 1
            continue

        if ln == "Table 2. User Requirements(Customer - Verified User)":
            _add_heading(doc, ln, 2)
            _add_table(doc, TABLE_2)
            i += 1
            while i < len(lines) and lines[i].strip() != "SYSTEM REQUIREMENTS SPECIFICATION":
                i += 1
            continue

        if ln == "Table 3. Functional Requirements (Admin)":
            _add_heading(doc, ln, 2)
            _add_table(doc, TABLE_3)
            i += 1
            while i < len(lines) and lines[i].strip() != "Table 4. Functional Requirements (Verified Customer)":
                i += 1
            continue

        if ln == "Table 4. Functional Requirements (Verified Customer)":
            _add_heading(doc, ln, 2)
            _add_table(doc, TABLE_4)
            i += 1
            while i < len(lines) and lines[i].strip() != "Table 5. Non-functional Requirements":
                i += 1
            continue

        if ln == "Table 5. Non-functional Requirements":
            _add_heading(doc, ln, 2)
            _add_table(doc, TABLE_5)
            i += 1
            # After NFR table, continue normally.
            continue

        # Headings
        if ln in H1_SET:
            _add_heading(doc, ln, 1)
            i += 1
            continue

        if ln in H2_SET:
            _add_heading(doc, ln, 2)
            i += 1
            continue

        # Heuristic: treat standalone Title Case lines as H2.
        if re.fullmatch(r"[A-Z][A-Za-z0-9’'()\- ]{2,}", ln) and (":" not in ln) and (len(ln.split()) <= 6):
            _add_heading(doc, ln, 2)
            i += 1
            continue

        # Regular paragraph
        doc.add_paragraph(ln)
        i += 1

    return doc


def main() -> None:
    if not MD_PATH.exists():
        raise SystemExit(f"Missing input file: {MD_PATH}")

    md_text = MD_PATH.read_text(encoding="utf-8")
    doc = build_docx(md_text)
    doc.save(str(DOCX_PATH))
    print(f"Wrote: {DOCX_PATH}")


if __name__ == "__main__":
    main()
